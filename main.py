import configparser
from dotenv import load_dotenv
from restapiwp import RestApiWP
from openaiClient import OpenaiClient
import logging
import os
from docx import Document

# Carga del fichero de configuracion
def loadConfig():
    # Crear un objeto ConfigParser
    config = configparser.ConfigParser()
    try:
        # Lee el archivo de configuración
        config.read('settings.ini')
    except Exception as err:
        logging.error(f"Error al cargar el fichero de configuraciones: {err}")
        raise
    return config

    
# Carga las contraseñas del fichero .env
def loadPasswords():
    try:
        load_dotenv()
    except Exception as err:
        logging.error(f"Error al cargar el fichero de contraseñas: {err}")
        raise
    return


#Genera el articulo solicitado
def generateArticle(config, keyword):

    # Formato del articulo
    new_post = {
        'title': 'Título del artículo',
        'content': 'Contenido del artículo',
        'status': 'draft',  # 'draft' para borradores, 'publish' para publicar
        'categories': [],   #  ['Cat1', 'Cat2'] o [] -> sin categorias
        'tags': [],         # ['Tag1', 'Tag2'] o [] -> sin tags
        'excerpt': 'Meta description'
    }
    client = OpenaiClient(config, logging)
    outline = client.generateOutline(keyword)
    title = client.generateTitle(keyword)
    content = client.generateContent(keyword, outline)
    meta_description = client.generate_meta_description(content)
    new_post['title'] = title
    new_post['excerpt'] = meta_description
    new_post['content'] = content
    return new_post

# Escribe el articulo a un fichero DOCX
def export_to_word(data):
    print("Generando articulo a .docx")
    filename = "./docx/" + data['title'].strip() + ".docx"
    try:
        if os.path.exists(filename):
            document = Document(filename)
        else:
            document = Document()

        document.add_heading(data['title'], level=1)
        document.add_paragraph(data['content'])
        document.add_paragraph(f"Meta Description: {data['excerpt']}")
        document.save(filename)
    except Exception as e:
        logging.error(f"Error en export_to_word(): {e}")
        raise


# Lee y devuelve en un array la lista de keywords(keywords.txt)
def readKeywordsFile(f_keywords):
    try:
        with open(f_keywords, 'r') as f:
            keywords = f.readlines()
        return keywords
    except FileNotFoundError:
        print(f"Error: El fichero '{f_keywords}' no se encontró.")
        logging.error(f"Error: El fichero '{f_keywords}' no se encontró.")
        return []
    except IOError as e:
        print(f"Error al leer el fichero '{f_keywords}': {e}")
        logging.error(f"Error al leer el fichero '{f_keywords}': {e}")
        return []

# Reescribe con las keywords que aun quedan por procesar en keywords.txt
def rewriteKeywordsFile(f_keywords, new_keywords):
    try:
        with open(f_keywords, 'w') as f:
            f.writelines(new_keywords)
    except IOError as e:
        print(f"Error al escribir en el fichero '{f_keywords}': {e}")

# Por cada linea de keywords.txt genera el articulo, hace una copia a un
# fichero.docx  escribe en la RESTAPI de  WP 
# y las keyword procesadas se llevan al fichero usedkeywords.txt
def initProcess():
    f_keywords = 'keywords.txt'
    f_used_keywords = 'usedkeywords.txt'
    # Leemos las keywords de keywords.txt
    keywords = readKeywordsFile(f_keywords)
    if not keywords:
        print("No hay keywords para procesar o no se pudo leer el fichero.")
        return
    # Lista para las keywords procesadas
    processedKeywords = []

    #Procesar cada Keyword
    for keyword in keywords:
        keyword = keyword.strip()
        if keyword:  # Ignorar líneas vacías
            try:
                # Procesando la keyword
                print(f"Procesando keyword: {keyword}")
                new_post = generateArticle(config, keyword) # Se genera el Articulo
                export_to_word(new_post)                    # Se genera el Articulo en .docx 
                response = RestApiWP(config, logging).writePost(new_post) # Se escribe el articulo en WP
                if response.status_code == 201:
                    print("Articulo generado insertado en WP correctamente ...")
                # Añadir la keyword procesada al fichero usedkeyword.txt
                with open(f_used_keywords, 'a') as used_file:
                    used_file.write(keyword + '\n')

                # Añadir la keyword a la lista de keywords procesadas
                processedKeywords.append(keyword + '\n')
            except IOError as e:
                print(f"Error al procesar la keyword '{keyword}': {e}")
            # Eliminar las líneas procesadas del fichero keywords.txt

        new_keywords = [kw for kw in keywords if kw not in processedKeywords]
        rewriteKeywordsFile(f_keywords, new_keywords)


if __name__ == "__main__":

    logging.basicConfig(filename="error.log", level=logging.ERROR, format="%(asctime)s:%(levelname)s:%(message)s")
    config = loadConfig()
    loadPasswords()
    initProcess()
    print("Proceso completado ....")
